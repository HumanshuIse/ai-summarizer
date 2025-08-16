# # backend/main.py
# import os
# from fastapi import FastAPI, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# from dotenv import load_dotenv
# import resend

# # --- MODIFICATION: Import Groq's library ---
# import groq

# # Load environment variables from the .env file
# load_dotenv()

# # Initialize the FastAPI app
# app = FastAPI()

# # --- CORS Middleware ---
# origins = [
#     "http://localhost:3000",
#     "http://localhost:5173",
# ]
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=origins,
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # --- Pydantic Models ---
# class SummarizeRequest(BaseModel):
#     transcript: str
#     prompt: str

# class ShareRequest(BaseModel):
#     summary: str
#     recipients: str

# # --- API Endpoints ---
# @app.get("/")
# def read_root():
#     return {"message": "AI Summarizer API is running"}

# @app.post("/api/summarize")
# async def generate_summary(request: SummarizeRequest):
#     # --- THIS FUNCTION IS NOW REVERTED TO USE GROQ ---
#     try:
#         client = groq.Groq(api_key=os.environ.get("GROQ_API_KEY"))
#         chat_completion = client.chat.completions.create(
#             messages=[
#                 {
#                     "role": "system",
#                     "content": "You are an expert meeting summarizer. Provide clear, structured summaries based on the user's instructions.",
#                 },
#                 {
#                     "role": "user",
#                     "content": f"Here is the transcript:\n\n{request.transcript}\n\nPlease apply this instruction: \"{request.prompt}\"",
#                 },
#             ],
#             model="llama3-8b-8192",
#         )
#         summary = chat_completion.choices[0].message.content
#         return {"summary": summary}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


# @app.post("/api/share")
# async def share_summary(request: ShareRequest):
#     try:
#         resend.api_key = os.environ.get("RESEND_API_KEY")
#         recipient_list = [email.strip() for email in request.recipients.split(',')]
        
#         params = {
#             "from": "Meeting Summarizer <onboarding@resend.dev>",
#             "to": recipient_list,
#             "subject": "Your AI-Generated Meeting Summary",
#             "html": f"<p>Here is the summary you requested:</p><pre>{request.summary}</pre>",
#         }
        
#         email = resend.Emails.send(params)
#         return {"message": "Email sent successfully!", "data": email}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
# backend/main.py
# import os
# from fastapi import FastAPI, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# from dotenv import load_dotenv
# import groq
# import resend
# from typing import Literal

# # --- NEW IMPORTS ---
# from io import BytesIO
# from docx import Document
# from reportlab.lib.pagesizes import letter
# from reportlab.platypus import SimpleDocTemplate, Paragraph
# from reportlab.lib.styles import getSampleStyleSheet

# # Load environment variables
# load_dotenv()

# # Initialize the FastAPI app
# app = FastAPI()

# # CORS Middleware
# origins = ["http://localhost:3000", "http://localhost:5173"]
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=origins,
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # --- Pydantic Models ---
# class SummarizeRequest(BaseModel):
#     transcript: str
#     prompt: str

# class ShareRequest(BaseModel):
#     summary: str
#     recipients: str
#     # --- NEW FIELD ---
#     # Add a format field. "plaintext" is the default.
#     format: Literal["plaintext", "pdf", "docx"] = "plaintext"

# # --- HELPER FUNCTIONS FOR FILE CREATION ---
# def create_docx_from_text(text: str) -> BytesIO:
#     document = Document()
#     document.add_paragraph(text)
#     file_stream = BytesIO()
#     document.save(file_stream)
#     file_stream.seek(0)
#     return file_stream

# def create_pdf_from_text(text: str) -> BytesIO:
#     file_stream = BytesIO()
#     doc = SimpleDocTemplate(file_stream, pagesize=letter)
#     styles = getSampleStyleSheet()
#     story = [Paragraph(line.replace('\n', '<br/>'), styles['Normal']) for line in text.split('\n')]
#     doc.build(story)
#     file_stream.seek(0)
#     return file_stream

# # --- API Endpoints ---
# @app.get("/")
# def read_root():
#     return {"message": "AI Summarizer API is running"}

# @app.post("/api/summarize")
# async def generate_summary(request: SummarizeRequest):
#     # (This endpoint remains unchanged)
#     try:
#         client = groq.Groq(api_key=os.environ.get("GROQ_API_KEY"))
#         chat_completion = client.chat.completions.create(
#             messages=[
#                 {"role": "system", "content": "You are an expert meeting summarizer..."},
#                 {"role": "user", "content": f"Here is the transcript:\n\n{request.transcript}\n\nPlease apply this instruction: \"{request.prompt}\""},
#             ],
#             model="llama3-8b-8192",
#         )
#         summary = chat_completion.choices[0].message.content
#         return {"summary": summary}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/api/share")
# async def share_summary(request: ShareRequest):
#     # --- THIS ENDPOINT IS NOW UPDATED ---
#     try:
#         resend.api_key = os.environ.get("RESEND_API_KEY")
#         recipient_list = [email.strip() for email in request.recipients.split(',')]
        
#         attachments = []
#         html_content = f"<p>Please find your meeting summary attached.</p>"

#         if request.format == "pdf":
#             pdf_bytes = create_pdf_from_text(request.summary)
#             attachments.append({
#                 "filename": "summary.pdf",
#                 "content": pdf_bytes.getvalue()
#             })
#         elif request.format == "docx":
#             docx_bytes = create_docx_from_text(request.summary)
#             attachments.append({
#                 "filename": "summary.docx",
#                 "content": docx_bytes.getvalue()
#             })
#         else: # Plaintext
#             html_content = f"<p>Here is the summary you requested:</p><pre>{request.summary}</pre>"

#         params = {
#             "from": "Meeting Summarizer <onboarding@resend.dev>",
#             "to": recipient_list,
#             "subject": "Your AI-Generated Meeting Summary",
#             "html": html_content,
#             "attachments": attachments
#         }
        
#         email = resend.Emails.send(params)
#         return {"message": "Email sent successfully!", "data": email}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
# backend/main.py
import os
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
import groq
import resend
from typing import Literal
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# Load environment variables
load_dotenv()

# Initialize the FastAPI app
app = FastAPI()

# CORS Middleware
origins = ["http://localhost:3000", "http://localhost:5173"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic Models
class SummarizeRequest(BaseModel):
    transcript: str
    prompt: str

class ShareRequest(BaseModel):
    summary: str
    recipients: str
    format: Literal["plaintext", "pdf", "docx"] = "plaintext"

# Helper functions for file creation
def create_docx_from_text(text: str) -> BytesIO:
    document = Document()
    document.add_paragraph(text)
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_pdf_from_text(text: str) -> BytesIO:
    file_stream = BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph(line.replace('\n', '<br/>'), styles['Normal']) for line in text.split('\n')]
    doc.build(story)
    file_stream.seek(0)
    return file_stream

# API Endpoints
@app.get("/")
def read_root():
    return {"message": "AI Summarizer API is running"}

@app.post("/api/summarize")
async def generate_summary(request: SummarizeRequest):
    try:
        client = groq.Groq(api_key=os.environ.get("GROQ_API_KEY"))
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are an expert meeting summarizer..."},
                {"role": "user", "content": f"Here is the transcript:\n\n{request.transcript}\n\nPlease apply this instruction: \"{request.prompt}\""},
            ],
            model="llama3-8b-8192",
        )
        summary = chat_completion.choices[0].message.content
        return {"summary": summary}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/share")
async def share_summary(request: ShareRequest):
    try:
        resend.api_key = os.environ.get("RESEND_API_KEY")
        recipient_list = [email.strip() for email in request.recipients.split(',')]
        
        attachments = []
        html_content = f"<p>Please find your meeting summary attached.</p>"

        if request.format == "pdf":
            pdf_bytes = create_pdf_from_text(request.summary)
            attachments.append({
                "filename": "summary.pdf",
                "content": pdf_bytes.getvalue()
            })
        elif request.format == "docx":
            docx_bytes = create_docx_from_text(request.summary)
            attachments.append({
                "filename": "summary.docx",
                "content": docx_bytes.getvalue()
            })
        else:
            html_content = f"<p>Here is the summary you requested:</p><pre>{request.summary}</pre>"

        params = {
            "from": "Meeting Summarizer <onboarding@resend.dev>",
            "to": recipient_list,
            "subject": "Your AI-Generated Meeting Summary",
            "html": html_content,
            "attachments": attachments
        }
        
        email = resend.Emails.send(params)
        
        # --- THIS IS THE CORRECTED LINE ---
        return {"message": "Email sent successfully!"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))